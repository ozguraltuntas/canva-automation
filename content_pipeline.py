"""content_pipeline.py — LLM (Claude / GPT-5.5) ile Amazon listing içeriği üretir.

Akış: standart_images + raw + Text edit.docx + user_text + OEM → LLM → parse → Content.docx
"""
import base64
import os
import re
from pathlib import Path

from dotenv import load_dotenv

load_dotenv()

CLAUDE_MODEL = "claude-opus-4-7"
OPENAI_MODEL = "gpt-5.5"

IMG_EXTS = {".jpg", ".jpeg", ".png", ".webp"}
MAX_IMAGES = 30  # safety cap

# LLM çıktısının sadece bu 4 bölümü dönmesi için ek talimat (kullanıcı görmez):
RESPONSE_OVERRIDE = (
    "\n\n---\n"
    "INPUT SOURCES — IMPORTANT:\n"
    "The user message may contain sections marked '## Category metadata' or "
    "'## Product-specific notes from user'. Treat the content of those sections as "
    "VERIFIED FACTS about the product. The 'use ONLY text visible in the image' rule "
    "applies to the IMAGES ONLY — it does NOT restrict the metadata/notes blocks. "
    "You MUST incorporate facts from those blocks into TITLE, BULLET POINTS, "
    "DESCRIPTION, and GENERIC KEYWORDS where relevant (brand, product type, color, "
    "fitment side, set/quantity, years-in-business claims, vehicle compatibility, etc.).\n"
    "\n"
    "RESPONSE FORMAT — STRICT:\n"
    "Return ONLY the four sections below, in this exact order, with these exact headers:\n"
    "TITLE:\n<title text>\n\n"
    "BULLET POINTS:\n<5 bullets, each on its own line>\n\n"
    "DESCRIPTION:\n<paragraphs>\n\n"
    "GENERIC KEYWORDS:\n<keywords>\n\n"
    "Do NOT include MISSING INFORMATION, FOLLOW-UP PROCESS, commentary, or any other section.\n"
    "Do NOT add lists describing what you read from the images.\n"
    "Output the four sections only."
)


# ---------------------------------------------------------------------------
# Text edit.docx
# ---------------------------------------------------------------------------

def read_text_edit_docx(path: Path) -> str:
    if not path.exists():
        return ""
    from docx import Document
    doc = Document(str(path))
    lines = [p.text for p in doc.paragraphs]
    return "\n".join(line for line in lines if line.strip())


# ---------------------------------------------------------------------------
# Images
# ---------------------------------------------------------------------------

def _list_images(folder: Path) -> list:
    if not folder.exists() or not folder.is_dir():
        return []
    return sorted(p for p in folder.iterdir() if p.suffix.lower() in IMG_EXTS)


def collect_images(category_dir: Path, parent_dir: Path) -> list:
    """standart_images/* (kategori ortak) + parent/raw/* (ürüne özel telefon foto)."""
    standart = _list_images(category_dir / "standart_images")
    raw = _list_images(parent_dir / "raw")
    combined = standart + raw
    if len(combined) > MAX_IMAGES:
        combined = combined[:MAX_IMAGES]
    return combined


def encode_image_b64(path: Path) -> tuple:
    """Returns (base64_str, media_type)."""
    suffix = path.suffix.lower().lstrip(".")
    media = "jpeg" if suffix == "jpg" else suffix
    media_type = f"image/{media}"
    data = base64.standard_b64encode(path.read_bytes()).decode("ascii")
    return data, media_type


# ---------------------------------------------------------------------------
# LLM calls
# ---------------------------------------------------------------------------

def _build_user_text_block(text_edit: str, user_text: str, oem: str) -> str:
    parts = []
    if text_edit.strip():
        parts.append(
            "## Category metadata (Text edit.docx) — AUTHORITATIVE, use these facts\n"
            "These lines come from a curated category metadata file. Treat them as VERIFIED "
            "FACTS about this product line and incorporate them into the listing where "
            "relevant (brand name, product type, color, fitment side, set/quantity, "
            "years-in-business claims, etc.). They are not restricted by the 'use only "
            "image text' rule.\n\n"
            f"{text_edit.strip()}"
        )
    if user_text.strip():
        parts.append(
            "## Product-specific notes from user — AUTHORITATIVE, use these facts\n"
            "Verified information provided by the seller about this specific SKU "
            "(typically vehicle compatibility, model years). Incorporate into the "
            "listing as appropriate.\n\n"
            f"{user_text.strip()}"
        )
    if oem.strip():
        parts.append(
            f"## OEM code(s)\n{oem.strip()}\n"
            "(These are reference codes — they will be appended to title/keywords by the program after your output. "
            "Do NOT include them yourself.)"
        )
    if not parts:
        return "Generate the listing using ONLY the visible text from the images."
    return "\n\n".join(parts)


def call_claude(prompt: str, images: list, user_text: str, text_edit: str,
                oem: str, model: str = CLAUDE_MODEL) -> str:
    import anthropic
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise RuntimeError("ANTHROPIC_API_KEY .env içinde tanımlı değil.")
    client = anthropic.Anthropic(api_key=api_key)

    content = []
    for img in images:
        b64, media_type = encode_image_b64(img)
        content.append({
            "type": "image",
            "source": {"type": "base64", "media_type": media_type, "data": b64},
        })
    content.append({"type": "text", "text": _build_user_text_block(text_edit, user_text, oem)})

    msg = client.messages.create(
        model=model,
        max_tokens=4096,
        system=prompt + RESPONSE_OVERRIDE,
        messages=[{"role": "user", "content": content}],
    )
    return "".join(block.text for block in msg.content if hasattr(block, "text"))


def call_openai(prompt: str, images: list, user_text: str, text_edit: str,
                oem: str, model: str = OPENAI_MODEL) -> str:
    import openai
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY .env içinde tanımlı değil.")
    client = openai.OpenAI(api_key=api_key)

    content = []
    for img in images:
        b64, media_type = encode_image_b64(img)
        content.append({
            "type": "image_url",
            "image_url": {"url": f"data:{media_type};base64,{b64}"},
        })
    content.append({"type": "text", "text": _build_user_text_block(text_edit, user_text, oem)})

    resp = client.chat.completions.create(
        model=model,
        max_completion_tokens=4096,
        messages=[
            {"role": "system", "content": prompt + RESPONSE_OVERRIDE},
            {"role": "user", "content": content},
        ],
    )
    return resp.choices[0].message.content or ""


# ---------------------------------------------------------------------------
# Parse + OEM + validate
# ---------------------------------------------------------------------------

SECTION_PATTERNS = {
    "title": r"^TITLE\s*:\s*",
    "bullets": r"^BULLET\s*POINTS\s*:\s*",
    "description": r"^DESCRIPTION\s*:\s*",
    "keywords": r"^GENERIC\s*KEYWORDS\s*:\s*",
}


def parse_llm_output(raw: str) -> dict:
    """LLM çıktısını TITLE / BULLETS / DESCRIPTION / KEYWORDS sözlüğüne ayır."""
    text = raw.strip()
    # Stop on these even if model still tries to add them:
    for stop in ("MISSING INFORMATION", "FOLLOW-UP PROCESS"):
        idx = text.find(stop)
        if idx >= 0:
            text = text[:idx].rstrip()

    lines = text.splitlines()
    sections = {"title": "", "bullets": "", "description": "", "keywords": ""}
    current = None
    buf: list = []

    def flush():
        if current is not None:
            sections[current] = "\n".join(buf).strip()

    for line in lines:
        matched = None
        for key, pat in SECTION_PATTERNS.items():
            if re.match(pat, line.strip(), flags=re.IGNORECASE):
                matched = key
                break
        if matched:
            flush()
            current = matched
            # If header line has trailing content after "TITLE:" keep it
            after = re.sub(SECTION_PATTERNS[matched], "", line.strip(), flags=re.IGNORECASE).strip()
            buf = [after] if after else []
        else:
            if current is not None:
                buf.append(line)
    flush()

    # Title typically single line — strip empties
    sections["title"] = sections["title"].strip()
    sections["keywords"] = sections["keywords"].strip()
    return sections


def apply_oem(parsed: dict, oem: str, targets: list) -> dict:
    """OEM kod(lar)ını title ve/veya keywords sonuna boşluk + OEM ekleyerek koy."""
    oem = (oem or "").strip()
    if not oem or not targets:
        return parsed
    out = dict(parsed)
    if "title" in targets and out.get("title"):
        out["title"] = f"{out['title']} {oem}".strip()
    if "keywords" in targets and out.get("keywords"):
        out["keywords"] = f"{out['keywords']} {oem}".strip()
    return out


def validate_lengths(parsed: dict) -> list:
    warnings = []
    title_len = len(parsed.get("title", ""))
    kw_len = len(parsed.get("keywords", ""))
    if title_len > 200:
        warnings.append(f"Title 200 karakteri aşıyor: {title_len} karakter")
    if kw_len > 250:
        warnings.append(f"Generic Keywords 250 karakteri aşıyor: {kw_len} karakter")
    return warnings


# ---------------------------------------------------------------------------
# Orchestrator + docx writer
# ---------------------------------------------------------------------------

def generate_content(provider: str, category_dir: Path, parent_dir: Path,
                     user_text: str, oem: str, oem_targets: list,
                     prompt_text: str) -> dict:
    """LLM'e gönder, parse+OEM+validate yap. Döner: dict with title/bullets/description/keywords/raw/warnings/images."""
    images = collect_images(category_dir, parent_dir)
    text_edit_path = category_dir / "Text edit.docx"
    text_edit = read_text_edit_docx(text_edit_path)

    if provider == "claude":
        raw = call_claude(prompt_text, images, user_text, text_edit, oem)
    elif provider in ("openai", "gpt"):
        raw = call_openai(prompt_text, images, user_text, text_edit, oem)
    else:
        raise ValueError(f"Unknown provider: {provider}")

    parsed = parse_llm_output(raw)
    parsed = apply_oem(parsed, oem, oem_targets)
    warnings = validate_lengths(parsed)
    parsed["raw"] = raw
    parsed["warnings"] = warnings
    parsed["image_count"] = len(images)
    parsed["images"] = [str(p) for p in images]
    parsed["text_edit"] = text_edit
    return parsed


def save_content_docx(parsed: dict, output_path: Path) -> None:
    """TITLE / BULLET POINTS / DESCRIPTION / GENERIC KEYWORDS bölümleriyle docx üret."""
    from docx import Document
    doc = Document()
    doc.add_paragraph("TITLE:")
    doc.add_paragraph(parsed.get("title", ""))
    doc.add_paragraph("BULLET POINTS:")
    bullets = parsed.get("bullets", "").strip()
    if bullets:
        for line in bullets.splitlines():
            line = line.strip()
            if line:
                doc.add_paragraph(line)
    doc.add_paragraph("DESCRIPTION:")
    desc = parsed.get("description", "").strip()
    for para in re.split(r"\n\s*\n", desc):
        para = para.strip()
        if para:
            doc.add_paragraph(para)
    doc.add_paragraph("GENERIC KEYWORDS:")
    doc.add_paragraph(parsed.get("keywords", ""))
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
